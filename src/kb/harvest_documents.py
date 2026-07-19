from __future__ import annotations

import argparse
import hashlib
import html
import json
import mimetypes
import re
import sys
import time
import zipfile
from collections import Counter
from dataclasses import dataclass, field
from datetime import datetime, timezone
from html.parser import HTMLParser
from pathlib import Path
from typing import Any
from urllib.error import HTTPError, URLError
from urllib.parse import unquote, urljoin, urlparse, urlunparse
from urllib.request import Request, urlopen


USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
    "AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/126.0 Safari/537.36 "
    "AIEngineerAssessmentBot/0.1"
)

DOCUMENT_EXTENSIONS = {
    ".pdf",
    ".doc",
    ".docx",
    ".xls",
    ".xlsx",
    ".csv",
    ".txt",
    ".rtf",
    ".zip",
}

DEFAULT_ALLOWED_DOMAINS = {
    "lendingkart.com",
    "www.lendingkart.com",
    "media.lendingkart.com",
    "lendingkartfinance.com",
    "www.lendingkartfinance.com",
}

LINK_KEYWORD_RE = re.compile(
    r"\b(download|downloads|pdf|docx?|xlsx?|csv|form|policy|disclosure|"
    r"schedule|charges|terms|condition|notice|report|certificate|statement)\b",
    re.IGNORECASE,
)

HIGH_RELEVANCE_RE = re.compile(
    r"(business[-\s]?loan|schedule[-\s]?of[-\s]?charges|interest[-\s]?rate|processing[-\s]?fee|"
    r"eligibility|repayment|cibil|grievance|ombudsman|fair[-\s]?practice|"
    r"privacy|personal[-\s]?data|terms[-\s]?and[-\s]?conditions|key[-\s]?fact|kfs|"
    r"loan[-\s]?fraud|moratorium|restructuring|rbi|customer[-\s]?care)",
    re.IGNORECASE,
)

ARCHIVE_ONLY_RE = re.compile(
    r"(annual[-\s]?general[-\s]?meeting|\bagm\b|\begm\b|shareholder|board[-\s]?meeting|"
    r"auditor|director|voting|scrutinizer|postal[-\s]?ballot|annual[-\s]?report|"
    r"annual[-\s]?return|\bmgt[-_\s]?7\b|mgt[-_\s]?7|form[-_\s]?mgt[-_\s]?7|"
    r"[_-]egm[_-]|csr[-_\s]?policy)",
    re.IGNORECASE,
)

MEDIUM_RELEVANCE_RE = re.compile(
    r"(regulatory|compliance|operational[-\s]?risk|disclosure|policy|security|fraud|"
    r"notice|lendingkart[-\s]?finance|ltpl)",
    re.IGNORECASE,
)

SPACE_RE = re.compile(r"[ \t\r\f\v]+")
SAFE_FILENAME_RE = re.compile(r"[^A-Za-z0-9._-]+")


@dataclass
class LinkHit:
    url: str
    tag: str
    attr: str
    attrs: dict[str, str]
    text: str = ""


@dataclass
class Candidate:
    document_id: str
    url: str
    domain: str
    extension: str
    allowed_domain: bool
    reasons: list[str] = field(default_factory=list)
    link_texts: list[str] = field(default_factory=list)
    source_pages: list[dict[str, str]] = field(default_factory=list)
    kb_relevance: str = "unknown"
    candidate_category: str = "unknown"
    download_status: str = "not_attempted"
    http_status: int | None = None
    content_type: str | None = None
    bytes_downloaded: int = 0
    downloaded_path: str | None = None
    extracted_text_status: str = "not_attempted"
    extracted_text_chars: int = 0
    cleaned_text_path: str | None = None
    error: str | None = None


class LinkCollector(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.links: list[LinkHit] = []
        self._anchor_stack: list[int] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        attrs_dict = {key.lower(): value or "" for key, value in attrs}
        tag = tag.lower()

        if tag == "a" and attrs_dict.get("href"):
            self.links.append(LinkHit(attrs_dict["href"], tag, "href", attrs_dict))
            self._anchor_stack.append(len(self.links) - 1)
            return

        for attr in ("href", "src", "data"):
            value = attrs_dict.get(attr)
            if value and tag in {"link", "iframe", "embed", "object", "source"}:
                self.links.append(LinkHit(value, tag, attr, attrs_dict))
                break

    def handle_data(self, data: str) -> None:
        if self._anchor_stack:
            idx = self._anchor_stack[-1]
            self.links[idx].text += data

    def handle_endtag(self, tag: str) -> None:
        if tag.lower() == "a" and self._anchor_stack:
            self._anchor_stack.pop()


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


def short_hash(text: str, length: int = 12) -> str:
    return hashlib.sha256(text.encode("utf-8", errors="ignore")).hexdigest()[:length]


def normalize_space(value: str) -> str:
    value = html.unescape(value or "")
    value = value.replace("\u00a0", " ")
    return SPACE_RE.sub(" ", value).strip()


def normalize_url(url: str) -> str:
    parsed = urlparse(url)
    return urlunparse((parsed.scheme, parsed.netloc.lower(), parsed.path, parsed.params, parsed.query, ""))


def extension_from_url(url: str) -> str:
    path = unquote(urlparse(url).path)
    suffix = Path(path).suffix.lower()
    if suffix in DOCUMENT_EXTENSIONS:
        return suffix
    return ""


def is_http_url(url: str) -> bool:
    parsed = urlparse(url)
    return parsed.scheme in {"http", "https"} and bool(parsed.netloc)


def is_allowed_domain(domain: str, allowed_domains: set[str]) -> bool:
    domain = domain.lower()
    for allowed in allowed_domains:
        allowed = allowed.lower()
        if domain == allowed or domain.endswith("." + allowed):
            return True
    return False


def candidate_reasons(url: str, link: LinkHit) -> list[str]:
    reasons: list[str] = []
    ext = extension_from_url(url)
    text_blob = " ".join([url, link.text, " ".join(link.attrs.keys()), " ".join(link.attrs.values())])

    if ext:
        reasons.append(f"document_extension:{ext}")
    if "download" in link.attrs:
        reasons.append("download_attribute")
    if LINK_KEYWORD_RE.search(text_blob):
        reasons.append("document_keyword")

    # Avoid adding normal webpage links just because their anchor says "view".
    if not ext and "download_attribute" not in reasons:
        return []
    return sorted(set(reasons))


def classify_candidate(url: str, link_texts: list[str], reasons: list[str]) -> tuple[str, str]:
    # Classify from source-facing text and URL only. Internal discovery reasons
    # include generic terms like "document_extension" and should not boost KB relevance.
    del reasons
    blob = " ".join([url, " ".join(link_texts)])
    lowered = blob.lower()

    if ARCHIVE_ONLY_RE.search(blob):
        return "archive_only", "corporate_archive"

    if re.search(r"(schedule[-\s]?of[-\s]?charges|interest|processing[-\s]?fee|charges)", lowered):
        return "high", "fees_and_charges"
    if re.search(r"(eligibility|cibil|repayment|business[-\s]?loan)", lowered):
        return "high", "loan_product"
    if re.search(r"(grievance|ombudsman|fair[-\s]?practice|privacy|personal[-\s]?data|fraud)", lowered):
        return "high", "customer_protection"
    if re.search(r"(co[-\s]?lending|partner|moratorium|restructuring|rbi|regulatory)", lowered):
        return "medium", "regulatory"
    if HIGH_RELEVANCE_RE.search(blob):
        return "high", "loan_or_compliance"
    if MEDIUM_RELEVANCE_RE.search(blob):
        return "medium", "compliance_archive"
    return "low", "unclassified_document"


def load_source_records(project_root: Path) -> list[dict[str, str]]:
    manifest_path = project_root / "data" / "kb" / "source_manifest.json"
    if manifest_path.exists():
        records = json.loads(manifest_path.read_text(encoding="utf-8"))
        output: list[dict[str, str]] = []
        for record in records:
            raw_file = record.get("raw_file") or record.get("raw_path")
            if raw_file:
                output.append(
                    {
                        "source_id": str(record.get("source_id") or Path(raw_file).stem),
                        "url": str(record.get("url") or ""),
                        "raw_file": str(raw_file),
                    }
                )
        if output:
            return output

    report_path = project_root / "data" / "kb" / "extraction_report.json"
    if report_path.exists():
        report = json.loads(report_path.read_text(encoding="utf-8"))
        output = []
        for record in report.get("results", []):
            raw_file = record.get("raw_path")
            if raw_file:
                output.append(
                    {
                        "source_id": str(record.get("source_id") or Path(raw_file).stem),
                        "url": str(record.get("url") or ""),
                        "raw_file": str(raw_file),
                    }
                )
        if output:
            return output

    sources_path = project_root / "data" / "sources.json"
    if sources_path.exists():
        sources = json.loads(sources_path.read_text(encoding="utf-8"))
        return [
            {
                "source_id": str(source["source_id"]),
                "url": str(source["url"]),
                "raw_file": str(project_root / "data" / "raw" / f"{source['source_id']}.html"),
            }
            for source in sources
        ]

    return []


def merge_unique(items: list[Any], value: Any, limit: int | None = None) -> None:
    if value in items:
        return
    if limit is not None and len(items) >= limit:
        return
    items.append(value)


def discover_candidates(
    project_root: Path,
    allowed_domains: set[str],
    include_external: bool,
) -> tuple[list[Candidate], list[Candidate], dict[str, Any]]:
    by_url: dict[str, Candidate] = {}
    skipped_external: dict[str, Candidate] = {}
    scanned_files = 0
    parsed_links = 0

    for source in load_source_records(project_root):
        raw_path = Path(source["raw_file"])
        if not raw_path.exists():
            continue

        scanned_files += 1
        try:
            html_text = raw_path.read_text(encoding="utf-8", errors="ignore")
        except OSError:
            continue

        collector = LinkCollector()
        collector.feed(html_text)
        parsed_links += len(collector.links)

        base_url = source.get("url") or raw_path.as_uri()
        for link in collector.links:
            href = normalize_space(link.url)
            if not href or href.startswith(("#", "mailto:", "tel:", "javascript:", "data:")):
                continue

            absolute_url = normalize_url(urljoin(base_url, href))
            if not is_http_url(absolute_url):
                continue

            reasons = candidate_reasons(absolute_url, link)
            if not reasons:
                continue

            domain = urlparse(absolute_url).netloc.lower()
            allowed = is_allowed_domain(domain, allowed_domains)
            target = by_url if allowed or include_external else skipped_external
            document_id = "doc_" + short_hash(absolute_url)
            candidate = target.get(absolute_url)
            if candidate is None:
                candidate = Candidate(
                    document_id=document_id,
                    url=absolute_url,
                    domain=domain,
                    extension=extension_from_url(absolute_url),
                    allowed_domain=allowed,
                )
                target[absolute_url] = candidate

            for reason in reasons:
                merge_unique(candidate.reasons, reason)
            link_text = normalize_space(link.text)
            if link_text:
                merge_unique(candidate.link_texts, link_text, limit=20)
            merge_unique(
                candidate.source_pages,
                {"source_id": source["source_id"], "url": source.get("url", "")},
            )

    all_candidates = list(by_url.values())
    external_candidates = list(skipped_external.values())
    for candidate in all_candidates + external_candidates:
        relevance, category = classify_candidate(candidate.url, candidate.link_texts, candidate.reasons)
        candidate.kb_relevance = relevance
        candidate.candidate_category = category
        if not candidate.allowed_domain:
            candidate.download_status = "skipped_external_domain"

    all_candidates.sort(key=lambda item: (item.kb_relevance, item.candidate_category, item.url))
    external_candidates.sort(key=lambda item: item.url)
    stats = {
        "scanned_raw_html_files": scanned_files,
        "parsed_links": parsed_links,
        "candidate_count": len(all_candidates),
        "skipped_external_candidate_count": len(external_candidates),
    }
    return all_candidates, external_candidates, stats


def safe_output_name(url: str, content_type: str | None = None) -> str:
    parsed = urlparse(url)
    basename = unquote(Path(parsed.path).name) or "downloaded_document"
    basename = basename.split("?")[0].strip(". ") or "downloaded_document"
    basename = SAFE_FILENAME_RE.sub("_", basename)
    suffix = Path(basename).suffix.lower()

    if not suffix:
        guessed = mimetypes.guess_extension((content_type or "").split(";")[0].strip())
        suffix = guessed if guessed in DOCUMENT_EXTENSIONS else ".bin"
        basename = basename + suffix

    stem = Path(basename).stem[:110].strip("._-") or "downloaded_document"
    return f"{stem}__{short_hash(url)}{Path(basename).suffix.lower()}"


def download_bytes(url: str, timeout: int, max_bytes: int) -> tuple[bytes, int | None, str | None]:
    request = Request(
        url,
        headers={
            "User-Agent": USER_AGENT,
            "Accept": "application/pdf,application/octet-stream,text/csv,text/plain,*/*;q=0.8",
            "Accept-Language": "en-IN,en;q=0.9",
        },
    )
    with urlopen(request, timeout=timeout) as response:
        status = getattr(response, "status", None)
        content_type = response.headers.get("Content-Type")
        body = response.read(max_bytes + 1)
        if len(body) > max_bytes:
            raise ValueError(f"download_exceeded_max_bytes:{max_bytes}")
        return body, status, content_type


def extract_pdf_text(path: Path) -> tuple[str, str]:
    pypdf_page_count: int | None = None
    pypdf_error: str | None = None

    try:
        from pypdf import PdfReader  # type: ignore

        reader = PdfReader(str(path))
        pypdf_page_count = len(reader.pages)
        parts = []
        for idx, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                parts.append(f"--- page {idx} ---\n{text.strip()}")
        if parts:
            return "\n\n".join(parts).strip(), f"pypdf_pages:{pypdf_page_count}"
    except ImportError:
        pass
    except Exception as exc:  # noqa: BLE001 - fallback parser may still succeed.
        pypdf_error = f"{type(exc).__name__}:{exc}"

    try:
        import pdfplumber  # type: ignore

        parts = []
        with pdfplumber.open(str(path)) as pdf:
            for idx, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                if text.strip():
                    parts.append(f"--- page {idx} ---\n{text.strip()}")
            page_count = len(pdf.pages)
        if parts:
            return "\n\n".join(parts).strip(), f"pdfplumber_pages:{page_count}"
        return "", f"pdf_text_empty_pages:{page_count}"
    except ImportError:
        if pypdf_page_count is not None:
            return "", f"pdf_text_empty_pages:{pypdf_page_count}"
        if pypdf_error:
            return "", f"pdf_extract_failed:{pypdf_error}"
        return "", "pdf_parser_unavailable"
    except Exception as exc:  # noqa: BLE001 - keep the exact failure in the report.
        if pypdf_page_count is not None:
            return "", f"pdf_text_empty_pages:{pypdf_page_count}"
        if pypdf_error:
            return "", f"pdf_extract_failed:{pypdf_error};pdfplumber:{type(exc).__name__}:{exc}"
        return "", f"pdf_extract_failed:{type(exc).__name__}:{exc}"


def extract_docx_text(path: Path) -> tuple[str, str]:
    try:
        import docx  # type: ignore
    except ImportError:
        return "", "docx_parser_unavailable"

    document = docx.Document(str(path))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = normalize_space(paragraph.text)
        if text:
            parts.append(text)
    for table_idx, table in enumerate(document.tables, start=1):
        parts.append(f"--- table {table_idx} ---")
        for row in table.rows:
            cells = [normalize_space(cell.text) for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip(), "python_docx"


def extract_spreadsheet_text(path: Path) -> tuple[str, str]:
    try:
        import openpyxl  # type: ignore
    except ImportError:
        return "", "spreadsheet_parser_unavailable"

    workbook = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        parts.append(f"--- sheet: {sheet.title} ---")
        row_count = 0
        for row in sheet.iter_rows(values_only=True):
            values = [normalize_space("" if cell is None else str(cell)) for cell in row]
            if any(values):
                parts.append(" | ".join(values))
            row_count += 1
            if row_count >= 5000:
                parts.append("--- truncated after 5000 rows ---")
                break
    return "\n".join(parts).strip(), "openpyxl"


def read_text_file(path: Path) -> tuple[str, str]:
    data = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return data.decode(encoding), f"text_decode:{encoding}"
        except UnicodeDecodeError:
            continue
    return "", "text_decode_failed"


def list_zip_contents(path: Path) -> tuple[str, str]:
    try:
        with zipfile.ZipFile(path) as archive:
            names = archive.namelist()
    except zipfile.BadZipFile:
        return "", "bad_zip_file"
    return "\n".join(names), f"zip_entries:{len(names)}"


def extract_document_text(path: Path) -> tuple[str, str]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf_text(path)
    if suffix == ".docx":
        return extract_docx_text(path)
    if suffix == ".xlsx":
        return extract_spreadsheet_text(path)
    if suffix in {".csv", ".txt", ".rtf"}:
        return read_text_file(path)
    if suffix == ".zip":
        return list_zip_contents(path)
    if suffix in {".doc", ".xls"}:
        return "", f"legacy_{suffix[1:]}_parser_unavailable"
    return "", "unsupported_file_type"


def write_cleaned_document(
    output_path: Path,
    candidate: Candidate,
    text: str,
    parser_status: str,
) -> None:
    header = [
        f"document_id: {candidate.document_id}",
        f"url: {candidate.url}",
        f"downloaded_path: {candidate.downloaded_path or ''}",
        f"content_type: {candidate.content_type or ''}",
        f"kb_relevance: {candidate.kb_relevance}",
        f"candidate_category: {candidate.candidate_category}",
        f"parser_status: {parser_status}",
        "source_pages: "
        + "; ".join(f"{page.get('source_id')} <{page.get('url')}>" for page in candidate.source_pages),
        f"extracted_at_utc: {utc_now()}",
        "",
        "--- CLEANED DOCUMENT TEXT ---",
        "",
    ]
    output_path.write_text("\n".join(header) + text.strip() + "\n", encoding="utf-8", errors="ignore")


def download_and_extract(
    candidates: list[Candidate],
    project_root: Path,
    timeout: int,
    max_mb: int,
    sleep_seconds: float,
    max_files: int | None,
) -> None:
    documents_dir = project_root / "data" / "raw" / "documents"
    cleaned_dir = project_root / "data" / "cleaned" / "documents"
    documents_dir.mkdir(parents=True, exist_ok=True)
    cleaned_dir.mkdir(parents=True, exist_ok=True)

    max_bytes = max_mb * 1024 * 1024
    attempted = 0

    for candidate in candidates:
        if not candidate.allowed_domain:
            candidate.download_status = "skipped_external_domain"
            continue
        if max_files is not None and attempted >= max_files:
            candidate.download_status = "skipped_max_files_limit"
            continue

        attempted += 1
        try:
            body, status, content_type = download_bytes(candidate.url, timeout=timeout, max_bytes=max_bytes)
            candidate.http_status = status
            candidate.content_type = content_type
            candidate.bytes_downloaded = len(body)
            filename = safe_output_name(candidate.url, content_type)
            output_path = documents_dir / filename
            output_path.write_bytes(body)
            candidate.downloaded_path = str(output_path)
            candidate.download_status = "downloaded"

            text, parser_status = extract_document_text(output_path)
            if text.strip():
                cleaned_path = cleaned_dir / f"{output_path.stem}.txt"
                write_cleaned_document(cleaned_path, candidate, text, parser_status)
                candidate.cleaned_text_path = str(cleaned_path)
                candidate.extracted_text_status = parser_status
                candidate.extracted_text_chars = len(text)
            else:
                candidate.extracted_text_status = parser_status
                candidate.extracted_text_chars = 0
        except HTTPError as exc:
            candidate.download_status = "failed"
            candidate.http_status = exc.code
            candidate.error = f"http_error:{exc.code}:{exc.reason}"
        except URLError as exc:
            candidate.download_status = "failed"
            candidate.error = f"url_error:{exc.reason}"
        except Exception as exc:  # noqa: BLE001 - report all source acquisition failures.
            candidate.download_status = "failed"
            candidate.error = f"{type(exc).__name__}:{exc}"

        if sleep_seconds:
            time.sleep(sleep_seconds)


def write_json(path: Path, data: Any) -> None:
    path.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")


def candidate_to_dict(candidate: Candidate) -> dict[str, Any]:
    return {
        "document_id": candidate.document_id,
        "url": candidate.url,
        "domain": candidate.domain,
        "extension": candidate.extension,
        "allowed_domain": candidate.allowed_domain,
        "reasons": candidate.reasons,
        "link_texts": candidate.link_texts,
        "source_pages": candidate.source_pages,
        "kb_relevance": candidate.kb_relevance,
        "candidate_category": candidate.candidate_category,
        "download_status": candidate.download_status,
        "http_status": candidate.http_status,
        "content_type": candidate.content_type,
        "bytes_downloaded": candidate.bytes_downloaded,
        "downloaded_path": candidate.downloaded_path,
        "extracted_text_status": candidate.extracted_text_status,
        "extracted_text_chars": candidate.extracted_text_chars,
        "cleaned_text_path": candidate.cleaned_text_path,
        "error": candidate.error,
    }


def write_candidate_markdown(
    path: Path,
    candidates: list[Candidate],
    external_candidates: list[Candidate],
    stats: dict[str, Any],
) -> None:
    lines = [
        "# Download Candidate Report",
        "",
        f"Generated at UTC: {utc_now()}",
        "",
        "## Discovery Stats",
        "",
        f"- Raw HTML files scanned: {stats['scanned_raw_html_files']}",
        f"- Links parsed: {stats['parsed_links']}",
        f"- Allowed-domain document candidates: {len(candidates)}",
        f"- External document candidates recorded but skipped: {len(external_candidates)}",
        "",
        "## Allowed Candidates",
        "",
        "| relevance | category | ext | source pages | url |",
        "| --- | --- | --- | ---: | --- |",
    ]
    for candidate in candidates:
        lines.append(
            "| {relevance} | {category} | {ext} | {pages} | {url} |".format(
                relevance=candidate.kb_relevance,
                category=candidate.candidate_category,
                ext=candidate.extension or "",
                pages=len(candidate.source_pages),
                url=candidate.url,
            )
        )

    if external_candidates:
        lines.extend(
            [
                "",
                "## External Candidates Skipped By Default",
                "",
                "| domain | ext | url |",
                "| --- | --- | --- |",
            ]
        )
        for candidate in external_candidates[:100]:
            lines.append(f"| {candidate.domain} | {candidate.extension or ''} | {candidate.url} |")

    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_download_markdown(path: Path, candidates: list[Candidate], stats: dict[str, Any]) -> None:
    download_counts = Counter(candidate.download_status for candidate in candidates)
    extraction_counts = Counter(candidate.extracted_text_status for candidate in candidates)
    relevance_counts = Counter(candidate.kb_relevance for candidate in candidates)
    total_bytes = sum(candidate.bytes_downloaded for candidate in candidates)

    lines = [
        "# Document Download Report",
        "",
        f"Generated at UTC: {utc_now()}",
        "",
        "## Summary",
        "",
        f"- Raw HTML files scanned: {stats['scanned_raw_html_files']}",
        f"- Allowed-domain document candidates: {len(candidates)}",
        f"- Downloaded bytes: {total_bytes}",
        "",
        "## Download Status",
        "",
    ]
    for status, count in sorted(download_counts.items()):
        lines.append(f"- {status}: {count}")

    lines.extend(["", "## Text Extraction Status", ""])
    for status, count in sorted(extraction_counts.items()):
        lines.append(f"- {status}: {count}")

    lines.extend(["", "## KB Relevance", ""])
    for relevance, count in sorted(relevance_counts.items()):
        lines.append(f"- {relevance}: {count}")

    lines.extend(
        [
            "",
            "## Downloaded Documents",
            "",
            "| status | text chars | relevance | category | bytes | url |",
            "| --- | ---: | --- | --- | ---: | --- |",
        ]
    )
    for candidate in candidates:
        lines.append(
            "| {status} | {chars} | {relevance} | {category} | {bytes} | {url} |".format(
                status=candidate.download_status,
                chars=candidate.extracted_text_chars,
                relevance=candidate.kb_relevance,
                category=candidate.candidate_category,
                bytes=candidate.bytes_downloaded,
                url=candidate.url,
            )
        )

    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Discover, download, and parse linked KB documents.")
    parser.add_argument("--project-root", type=Path, default=Path.cwd(), help="Project root containing data/raw.")
    parser.add_argument(
        "--allowed-domain",
        action="append",
        default=[],
        help="Domain allowed for downloads. Can be passed multiple times.",
    )
    parser.add_argument(
        "--include-external",
        action="store_true",
        help="Also download document links outside the allowed domains.",
    )
    parser.add_argument("--discover-only", action="store_true", help="Only write candidate reports; do not download.")
    parser.add_argument("--timeout", type=int, default=35, help="Per-document download timeout in seconds.")
    parser.add_argument("--max-mb", type=int, default=50, help="Maximum bytes per downloaded document in MB.")
    parser.add_argument("--sleep", type=float, default=0.4, help="Polite delay between downloads.")
    parser.add_argument("--max-files", type=int, default=None, help="Optional maximum number of downloads.")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    project_root = args.project_root.resolve()
    kb_dir = project_root / "data" / "kb"
    kb_dir.mkdir(parents=True, exist_ok=True)

    allowed_domains = set(DEFAULT_ALLOWED_DOMAINS)
    allowed_domains.update(domain.lower() for domain in args.allowed_domain)

    candidates, external_candidates, stats = discover_candidates(
        project_root=project_root,
        allowed_domains=allowed_domains,
        include_external=args.include_external,
    )

    write_json(kb_dir / "download_candidates.json", [candidate_to_dict(candidate) for candidate in candidates])
    write_json(
        kb_dir / "download_candidates_external_skipped.json",
        [candidate_to_dict(candidate) for candidate in external_candidates],
    )
    write_candidate_markdown(kb_dir / "download_candidates.md", candidates, external_candidates, stats)

    if not args.discover_only:
        download_and_extract(
            candidates=candidates,
            project_root=project_root,
            timeout=args.timeout,
            max_mb=args.max_mb,
            sleep_seconds=args.sleep,
            max_files=args.max_files,
        )

    report = {
        "generated_at_utc": utc_now(),
        "project_root": str(project_root),
        "discover_only": args.discover_only,
        "include_external": args.include_external,
        "allowed_domains": sorted(allowed_domains),
        "stats": stats,
        "download_status_counts": dict(Counter(candidate.download_status for candidate in candidates)),
        "extracted_text_status_counts": dict(Counter(candidate.extracted_text_status for candidate in candidates)),
        "kb_relevance_counts": dict(Counter(candidate.kb_relevance for candidate in candidates)),
        "candidates": [candidate_to_dict(candidate) for candidate in candidates],
        "external_candidates_skipped": [candidate_to_dict(candidate) for candidate in external_candidates],
    }
    write_json(kb_dir / "download_report.json", report)
    write_download_markdown(kb_dir / "download_report.md", candidates, stats)

    print(
        "discovered={0} external_skipped={1} downloaded={2} extracted={3}".format(
            len(candidates),
            len(external_candidates),
            sum(1 for candidate in candidates if candidate.download_status == "downloaded"),
            sum(1 for candidate in candidates if candidate.extracted_text_chars > 0),
        )
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
